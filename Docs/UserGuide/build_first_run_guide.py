from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "时光记-首次使用指南.docx"
SCREENSHOT = ROOT / "Assets" / "FirstRun" / "01-home.jpg"
FONT_NAME = "Hiragino Sans GB"


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def add_title(document, text, subtitle=None):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(28, 74, 129)
    paragraph.paragraph_format.space_after = Pt(6)
    if subtitle:
        sub = document.add_paragraph(subtitle)
        sub.style = document.styles["Subtitle"]


def add_heading(document, text, level=1):
    return document.add_heading(text, level=level)


def add_body(document, text, bold_prefix=None):
    paragraph = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    return paragraph


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_steps(document, items):
    for item in items:
        document.add_paragraph(item, style="List Number")


def build():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, color in [
        ("Title", 26, "1C4A81"),
        ("Subtitle", 11, "667085"),
        ("Heading 1", 17, "1C4A81"),
        ("Heading 2", 13, "2F6FAB"),
        ("Heading 3", 11, "344054"),
    ]:
        style = styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)

    add_title(document, "时光记首次使用指南", "给第一次打开时光记的你｜宝宝成长照片从此看得懂时间")
    add_body(document, "手机里的宝宝照片越来越多，过一段时间再打开，我们常常记得那个表情，却想不起照片是哪天拍的、宝宝当时多大。")
    add_body(document, "时光记读取照片原本的拍摄时间，结合宝宝生日，把年龄和记忆文字呈现在一张新图片上。")

    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    labels = [("本地处理", "照片不上传"), ("保留原图", "生成新图片"), ("自动计算", "拍摄时多大")]
    for cell, (title, detail) in zip(table.rows[0].cells, labels):
        cell.width = Inches(2.1)
        set_cell_shading(cell, "EAF4FF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title + "\n")
        r.bold = True
        r.font.color.rgb = RGBColor(20, 111, 199)
        p.add_run(detail)

    add_heading(document, "第一次打开：先完成一次设置")
    add_body(document, "真实首页截图已随指南资产单独提供；正式发布版将在补齐分享、任务和回存页面后统一排入图文手册。")

    add_heading(document, "1. 填写宝宝资料", 2)
    add_body(document, "打开“记忆对象”，填写宝宝昵称、你与宝宝的关系和出生日期。出生日期用于在本机计算每张照片拍摄时宝宝多大。")
    add_heading(document, "2. 保存第一套配置", 2)
    add_body(document, "进入“配置中心”，第一次可以先使用推荐显示方式。确认顶部真实记忆卡预览后，点“保存当前配置”。以后想换文字、Logo 标识或表达语气时再回来调整。")

    add_heading(document, "第一次处理照片")
    add_body(document, "推荐日常路径：Apple Photos → 分享 → 时光记 → 处理 → 通知 → Apple Photos")
    add_steps(document, [
        "打开 iPhone 的“照片”，选择想处理的宝宝照片。",
        "点“分享”，在分享面板中选择“时光记”。",
        "确认接收后等待本地处理。",
        "收到完成通知后，回到“照片”查看新生成的图片。",
    ])
    add_body(document, "每次最多选择 20 张。照片较多时请分批分享，处理和回存会更稳定。")

    add_heading(document, "以后怎么使用")
    add_body(document, "完成一次配置后，大多数时候不用先打开时光记。直接在 Apple Photos 选择照片并分享给时光记即可。")
    add_bullets(document, [
        "换宝宝或其他记录对象时，回到“记忆对象”。",
        "修改生日或纪念日时，回到“时间锚点”。",
        "更换预设、Logo 标识或文字时，回到“配置中心”。",
        "查看最近处理时，打开“任务”。",
    ])

    add_heading(document, "时间锚点：一句话看懂")
    add_body(document, "时间锚点是一个有意义的日期。宝宝生日是最常用的锚点：照片拍摄日期减去宝宝生日，就得到拍摄时的年龄。")
    add_body(document, "示例：主体“小满” + 生日锚点 + 照片拍摄日期 → 1岁2个月18天。")
    add_bullets(document, [
        "自然：今天小满1岁2个月18天",
        "成长：小满长到1岁2个月18天了",
        "温馨：陪小满走到1岁2个月18天",
        "极简：小满｜1岁2个月18天",
    ])
    add_body(document, "完整的五类锚点、25 种语气和变量公式，请查看同目录的《时间锚点与表达语气说明》。")

    add_heading(document, "常见问题")
    qa = [
        ("找不到时光记", "在分享面板的 App 列表滑到最后，点“更多”，找到时光记并启用。首次分享前请先保存一次当前配置。"),
        ("没有看到新图片", "打开“任务”查看状态，并确认已允许读取照片和保存结果。"),
        ("原图会被覆盖吗", "不会。时光记生成新文件，原始照片保持不变。"),
        ("一次可以处理多少张", "每次最多 20 张，较多照片请分批处理。"),
        ("Live Photo、RAW/DNG", "相关路径仍持续进行真机兼容性验证；动态效果与格式保留取决于输入和当前输出配置。"),
        ("删除 App", "已经保存到相册的图片仍在，但 App 本机容器中的宝宝资料、锚点、配置和任务记录可能被删除。"),
    ]
    for question, answer in qa:
        add_body(document, f"{question}：{answer}", bold_prefix=f"{question}：")

    add_heading(document, "隐私与反馈")
    add_bullets(document, [
        "照片处理在设备本地完成，不上传原始照片。",
        "TestFlight 反馈适合提交闪退、截图和录屏。",
        "邮件：serydoo@gmail.com",
        "小红书：ID 49956456623",
    ])

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("时光记 · 本地优先的记忆呈现引擎")
    document.save(OUTPUT)


if __name__ == "__main__":
    build()
